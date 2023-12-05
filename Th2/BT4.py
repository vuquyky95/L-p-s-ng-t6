import tkinter as tk
from tkinter import END, Scrollbar, Text, messagebox
import matplotlib.pyplot as plt
import pandas as pd
from numpy import array
import numpy as np
import os
from docx import Document


def export_text_to_word(text_widget, file_path):
    # Lấy nội dung từ ô văn bản
    content = text_widget.get("1.0", "end-1c")

    # Tạo một tài liệu Word mới
    doc = Document()

    # Thêm nội dung vào tài liệu
    doc.add_paragraph(content)

    charts = ['chart1.png', 'chart.png', 'chart2.png']

    for chart in charts:
        if os.path.exists(chart):
            doc.add_picture(chart)
    # Lưu tài liệu vào file Word
    doc.save(file_path)


def draw_chart():
    # code vẽ biểu đồ

    plt.savefig('chart.png')
    plt.savefig('chart2.png')
    plt.savefig('chart1.png')


def xuat_word():
    file_path = "output.docx"

    export_text_to_word(result_text, file_path)

    messagebox.showinfo("Thông báo", "Xuất file thành công")


df = pd.read_csv('diemPython.csv', index_col=0, header=0)
in_data = array(df.iloc[:, :])


def hienthidanhsach():
    result_text.insert(END, "Danh sách sinh viên:\n")
    result_text.insert(END, df.to_string(index=False) + "\n\n")


def tongsinhvien():
    sv = in_data[:, 1]
    tongsv = np.sum(sv)
    result_text.insert(END, "Tổng số sinh viên đi thi: " + str(tongsv) + " sinh viên\n")

    svtruot = in_data[:, 10]
    tongsvtruot = np.sum(svtruot)
    tongsvdat = tongsv - tongsvtruot

    percent_dat = (tongsvdat / tongsv) * 100
    percent_truot = (tongsvtruot / tongsv) * 100

    result_text.insert(END, "Tổng số sinh viên qua môn: " + str(tongsvdat) + " sinh viên (" + str(percent_dat) + "%)\n")
    result_text.insert(END, "Tổng số sinh viên trượt môn: " + str(tongsvtruot) + " sinh viên (" + str(
        percent_truot) + "%)\n\n")
    categories = ['1', '2', '3', '4', '5', '6', '7', '8', '9']
    values1 = np.sum(in_data[0:9, 2:10], axis=1).flatten()
    values2 = in_data[:, 10]
    total_values = values1 + values2

    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(18, 6))

    ax1.bar(categories, values1, color='green', label="Sinh viên đạt")
    ax1.bar(categories, values2, bottom=values1, color='red', label="Sinh viên trượt")

    for i, (v1, v2, tv) in enumerate(zip(values1, values2, total_values)):
        ax1.text(i, v1 / 2, str(v1), ha='center', va='bottom', color='white', fontweight='bold')
        ax1.text(i, v1 + v2 / 2, str(v2), ha='center', va='bottom', color='white', fontweight='bold')
        ax1.text(i, v1 + v2, str(tv), ha='center', va='bottom', color='black', fontweight='bold')

    ax1.set_title('Biểu đồ số sinh viên trong các lớp')
    ax1.set_ylabel('Số sinh viên')
    ax1.set_xlabel('Lớp')
    ax1.legend(loc='upper right')

    ax2.pie(total_values, labels=categories, autopct='%1.1f%%')
    ax2.set_title('Biểu đồ tỷ lệ sinh viên theo lớp')

    ax3.pie([percent_dat, percent_truot], labels=["Sinh viên đạt", "Sinh viên trượt"], autopct='%1.1f%%')
    ax3.set_title('Biểu đồ tỷ lệ sinh viên đạt và trượt')

    # Lưu hình ảnh
    plt.savefig('chart.png')
    plt.subplots_adjust(wspace=0.5)
    plt.show()


def sosinhvientruot():
    svtruot = in_data[:, 10]
    tongsvtruot = np.sum(svtruot)
    result_text.insert(END, "Tổng số sinh viên trượt môn:")
    result_text.insert(END, str(tongsvtruot) + "\n\n")

    categories1 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
    values1 = in_data[:, 10]

    plt.figure(2)
    plt.bar(categories1, values1, color='red', label="Sinh viên trượt")
    plt.title('Biểu đồ số sinh viên trượt của các lớp')
    plt.ylabel('Số sinh viên')
    plt.legend(loc='upper right')
    plt.show()


def sosinhviendat():
    sv = in_data[:, 1]
    tongsv = np.sum(sv)
    svtruot = in_data[:, 10]
    tongsvtruot = np.sum(svtruot)
    tongsvdat = tongsv - tongsvtruot
    result_text.insert(END, "Tổng số sinh viên qua môn: " + str(tongsvdat) + " sinh viên \n")

    categories2 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
    values2 = np.sum(in_data[0:9, 2:10], axis=1).flatten()

    plt.figure(3)
    plt.bar(categories2, values2, color='green', label="Sinh viên đạt")
    plt.title('Biểu đồ số sinh viên đạt của các lớp')
    plt.ylabel('Số sinh viên')
    plt.legend(loc='upper right')
    plt.show()


def sinhvienA():
    diemA = in_data[:, 3]
    maxa = diemA.max()
    i = np.argmax(diemA)
    result_text.insert(END, 'Lớp có nhiều sinh viên điểm A là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))

    categories1 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
    values1 = in_data[:, 3]

    plt.figure(4)
    bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm A")
    plt.title('Biểu đồ số sinh viên đạt điểm A của các lớp')
    plt.ylabel('Số sinh viên')
    plt.legend(loc='upper right')

    for bar, value in zip(bars, values1):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')

    plt.show()


def phodiem():
    labels = ["A", "B+", "B", "C+", "C", "D+", "D", "F"]
    colors = ['red', 'green', 'yellow', 'blue', 'orange', 'purple', 'pink', 'brown']
    diems = in_data[:, 3:11]
    total_students = diems.shape[0]
    grade_counts = diems.sum(axis=0)

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))

    # Bar chart
    ax1.bar(labels, grade_counts, color=colors)
    ax1.set_xlabel('Điểm')
    ax1.set_ylabel('Số lượng sinh viên')
    ax1.set_title('Phổ điểm')
    for i, v in enumerate(grade_counts):
        ax1.annotate(str(v), (i, v), ha='center', va='bottom')

    # Pie chart
    ax2.pie(grade_counts, labels=labels, colors=colors, autopct='%1.1f%%')
    ax2.axis('equal')
    ax2.set_title('Phổ điểm')
    # Lưu hình ảnh
    plt.savefig('chart1.png')
    plt.tight_layout()
    plt.show()


def reset_result():
    result_text.delete(1.0, END)
    plt.close('all')


def number_of_students_by_grade():
    labels = ["A", "B+", "B", "C+", "C", "D+", "D"]
    grade_counts = []
    for label in labels:
        grade_column = in_data[:, labels.index(label) + 3]
        grade_count = len(grade_column[grade_column == 1])
        grade_counts.append(grade_count)
    # Create a table in the result_text widget
    result_text.insert(END, "Số lượng sinh viên theo điểm:\n")
    result_text.insert(END, "--------------------------------\n")
    result_text.insert(END, "Điểm\tSố lượng sinh viên\n")
    result_text.insert(END, "--------------------------------\n")

    for label, count in zip(labels, grade_counts):
        result_text.insert(END, f"{label}\t{count}\n")
    result_text.insert(END, "--------------------------------\n")


# Tạo cửa sổ giao diện
window = tk.Tk()
window.title("Ứng dụng báo cáo")

# Tạo các nút chức năng
htdanhsach_button = tk.Button(window, text="Hiển thị danh sách", command=hienthidanhsach)
tongsinhvien_button = tk.Button(window, text="Tổng sinh viên đi thi", command=tongsinhvien)
sinhvientruot_button = tk.Button(window, text="Số sinh viên trượt", command=sosinhvientruot)
sinhviendat_button = tk.Button(window, text="Số sinh viên đạt", command=sosinhviendat)
sinhvienA_button = tk.Button(window, text="Sinh viên đạt điểm A", command=sinhvienA)
phodiem_button = tk.Button(window, text="Phổ điểm", command=phodiem)
reset_button = tk.Button(window, text="Reset", command=reset_result)
xuat_word_button = tk.Button(window, text="Xuất file word", command=xuat_word)
# Tạo ô văn bản để hiển thị kết quả
result_text = Text(window, wrap="word", height=20, width=200)
result_text_scrollbar = Scrollbar(window, command=result_text.yview)
result_text.config(yscrollcommand=result_text_scrollbar.set)

# Đặt vị trí các phần tử trên cửa sổ
htdanhsach_button.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
tongsinhvien_button.grid(row=1, column=0, padx=10, pady=10, sticky="ew")
sinhvientruot_button.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
sinhviendat_button.grid(row=3, column=0, padx=10, pady=10, sticky="ew")
sinhvienA_button.grid(row=4, column=0, padx=10, pady=10, sticky="ew")
phodiem_button.grid(row=5, column=0, padx=10, pady=10, sticky="ew")
reset_button.grid(row=6, column=0, padx=10, pady=10, sticky="ew")
result_text.grid(row=0, column=1, rowspan=7, padx=10, pady=10, sticky="nsew")
result_text_scrollbar.grid(row=0, column=2, rowspan=7, sticky="ns")
xuat_word_button.grid(row=7, column=0, padx=10, pady=10, sticky="ew")

number_of_students_button = tk.Button(window, text="Số lượng sinh viên theo điểm", command=number_of_students_by_grade)
number_of_students_button.grid(row=9, column=0, padx=10, pady=10, sticky="ew")

# Chạy ứng dụng
window.mainloop()